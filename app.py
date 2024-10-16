import docx

# Function to load the DOCX template
def load_template(docx_template_path):
    return docx.Document(docx_template_path)

# Function to replace placeholders across paragraphs and runs, ensuring placeholder replacement while preserving formatting
def replace_placeholders(doc, placeholders):
    for para in doc.paragraphs:
        # Only process paragraphs that contain placeholders
        if any(key in para.text for key in placeholders):
            # Combine the text from all runs in the paragraph
            full_text = "".join(run.text for run in para.runs)
            
            # Replace placeholders in the full text of the paragraph
            for key, value in placeholders.items():
                if key in full_text:
                    full_text = full_text.replace(key, value)

            # Reset the original runs, but preserve formatting for each run
            remaining_text = full_text
            for run in para.runs:
                run_length = len(run.text)
                run_format = run.font  # Save the current formatting (size, font, etc.)
                run.text = remaining_text[:run_length]  # Set the text for this run
                remaining_text = remaining_text[run_length:]  # Update the remaining text for subsequent runs
                
                # Apply original formatting back to the new text
                run.font.name = run_format.name
                run.font.size = run_format.size

            # If there is any text left after iterating through the runs, add it to a new run with the same formatting
            if remaining_text:
                new_run = para.add_run(remaining_text)
                new_run.font.name = run_format.name
                new_run.font.size = run_format.size
    
    return doc

# Function to save the updated DOCX file
def save_customized_docx(doc, output_path):
    doc.save(output_path)
